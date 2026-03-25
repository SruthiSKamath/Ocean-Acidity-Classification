import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_contribution_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Data Engineering Team Contributions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Cohort 23 — Ocean Acidity Level Categorization (SDG 14)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('\nThis document outlines the specific data engineering activities and preprocessing tasks performed by each member of the four-person data engineering team. All members contributed equally to the development of the data preprocessing pipeline.\n')
    
    # --- Team Member 1: Vaastav L Sanghvi ---
    h1 = doc.add_heading('1. Vaastav L Sanghvi — Dataset Discovery & Ingestion', level=2)
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run('Dataset Sourcing: ').bold = True
    p1.add_run('Identified, evaluated, and sourced the Surface Ocean CO₂ Atlas (SOCAT) v2022 dataset from the internet, which serves as the foundational data for this project.')
    p1_2 = doc.add_paragraph(style='List Bullet')
    p1_2.add_run('Data Loading Pipeline: ').bold = True
    p1_2.add_run('Designed the initial strategy to handle the large-scale 1.83 GB Parquet file, ensuring efficient column selection and memory management during the ingestion phase.')
    
    # --- Team Member 2: Saayanth ---
    h2 = doc.add_heading('2. Saayanth — Data Quality Assessment & Cleaning', level=2)
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('Quality Assessment: ').bold = True
    p2.add_run('Analyzed the dataset for missing values, invalid entries, and examined the distribution of quality control (QC) flags across millions of rows.')
    p2_2 = doc.add_paragraph(style='List Bullet')
    p2_2.add_run('Data Cleaning Implementation: ').bold = True
    p2_2.add_run('Implemented the filtering logic to retain only high-quality fCO₂ measurements (fCO2rec_flag == 2) and acceptable overall QC flags (A, B, D).')
    p2_3 = doc.add_paragraph(style='List Bullet')
    p2_3.add_run('Imputation Strategy: ').bold = True
    p2_3.add_run('Developed and applied the median-based imputation pipeline for missing numeric oceanographic variables to ensure robustness against outliers.')
    
    # --- Team Member 3: Sadiya Kulsum ---
    h3 = doc.add_heading('3. Sadiya Kulsum — Sensor Calibration & Feature Engineering', level=2)
    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run('Calibration Offset Tuning: ').bold = True
    p3.add_run('Designed the sensor calibration offset algorithm within the preprocessing pipeline. Conducted sensitivity analysis (from −15 to +15 µatm) to demonstrate how sensor biases affect the final acidity class distributions.')
    p3_2 = doc.add_paragraph(style='List Bullet')
    p3_2.add_run('Feature Engineering: ').bold = True
    p3_2.add_run('Engineered critical derived features, including the shipping traffic proxies (inverse distance to land, costal flags), seasonal temporal features (cyclical sine/cosine encoding), and oceanographic interaction metrics (e.g., SST-Salinity interaction).')
    
    # --- Team Member 4: Vrushank Skanda B ---
    h4 = doc.add_heading('4. Vrushank Skanda B — Target Extraction, EDA & Final Pipeline', level=2)
    p4 = doc.add_paragraph(style='List Bullet')
    p4.add_run('Target Variable Creation: ').bold = True
    p4.add_run('Defined the CO₂ thresholds for acidity classification (Safe < 380 µatm, Vulnerable 380-450 µatm, Critical > 450 µatm) and built the labeling logic.')
    p4_2 = doc.add_paragraph(style='List Bullet')
    p4_2.add_run('Outlier Treatment & EDA: ').bold = True
    p4_2.add_run('Performed Exploratory Data Analysis (EDA) plotting distributions and feature correlations. Implemented IQR-based winsorization to cap extreme outlier values without deleting valid samples.')
    p4_3 = doc.add_paragraph(style='List Bullet')
    p4_3.add_run('Scaling & Export: ').bold = True
    p4_3.add_run('Finalized the preprocessing by applying StandardScaler to features, encoding target labels, and exporting the final model-ready datasets (CSV and Parquet formats).')
    
    # Save document
    output_path = r'c:\Users\sejal\OneDrive\Desktop\datasets\ocean_acidity_project\Team_Contributions_Report.docx'
    doc.save(output_path)
    print(f"Successfully saved Word document to: {output_path}")

if __name__ == '__main__':
    create_contribution_doc()
